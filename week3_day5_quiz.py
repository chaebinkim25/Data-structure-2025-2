# -*- coding: utf-8 -*-
# 미니퀴즈 (주관식·1페이지) DOCX 생성 스크립트 + 수업 분량 설문
# 필요: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def set_kr(run, size=11, bold=False, italic=False, font_name="Malgun Gothic"):
    try:
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_line(doc, width_chars=60):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("─" * width_chars)
    set_kr(run, size=9)

def add_short(doc, num, stem, lines=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"{num}. {stem}")
    set_kr(run, size=11)
    for _ in range(lines):
        add_line(doc)

def build_one_page_quiz(output_path="doc/자료구조_3주차_5회차_미니퀴즈.docx"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("미니퀴즈 (주관식) — 스택")
    set_kr(run, size=14, bold=True)

    # 부제
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"날짜: {datetime.now().strftime('%Y-%m-%d')}   이름: ____________   학번: ____________________")
    set_kr(run, size=10, italic=True)

    # 안내
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("※ 간단명료하게 쓰세요. 풀이가 필요한 경우 핵심만 적기.")
    set_kr(run, size=9)

    # 문항
    q = 1
    add_short(doc, q, "스택(Stack)의 두 가지 기본 연산을 쓰고, 각각을 한 줄로 설명하시오.", lines=2); q+=1
    add_short(doc, q, "단방향 링크드 리스트에서 스택을 구현할 때, top을 어디로 두는 것이 유리한지와 그 이유를 쓰시오.", lines=2); q+=1
    add_short(doc, q, "POP 연산에서 발생할 수 있는 에러(예외 상황)를 쓰고, 그 원인을 설명하시오.", lines=2); q+=1
    add_short(doc, q, "배열로 구현한 스택에서 top == 1일 때와 top == capacity + 1일 때 각각 무엇을 의미하는지 쓰시오.", lines=2); q+=1
    add_short(doc, q, "배열로 구현한 스택에서 PUSH와 POP 동작을 설명하시오.", lines=2); q+=1
    add_short(doc, q, "미로 탐색에서 방문한 위치와 앞으로 갈 곳을 표시하는 이유를 간단히 설명하시오.", lines=2); q+=1
    add_short(doc, q, "미로 탐색 과정에서 PROCESS_NEIGHBORS 함수의 핵심 역할을 한두 줄로 설명하시오.", lines=2); q+=1

    # ★ 설문 문항 추가
    doc.add_paragraph()  # 빈 줄
    run = doc.add_paragraph().add_run("※ 수업 피드백")
    set_kr(run, size=11, bold=True)

    add_short(doc, "S1", "오늘 수업 분량은 적절했습니까? (너무 적다 / 적절하다 / 많다 중 선택하고 간단히 이유)", lines=2)
    add_short(doc, "S2", "추가적으로 원하는 설명이나 개선 사항이 있으면 간단히 적어주세요.", lines=2)

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    path = build_one_page_quiz()
    print("Saved:", path)
