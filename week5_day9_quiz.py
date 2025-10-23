# -*- coding: utf-8 -*-
# 미니퀴즈 (주관식·1페이지) DOCX 생성 스크립트 + 수업 분량 설문
# 필요: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def set_kr(run, size=11, bold=False, italic=False, font_name="Malgun Gothic"):
    try:
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_line(doc, width_chars=60):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("─" * width_chars)
    set_kr(run, size=9)

def add_short(doc, num, stem, lines=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"{num}. {stem}")
    set_kr(run, size=11)
    for _ in range(lines):
        add_line(doc)

def build_one_page_quiz(output_path="doc/자료구조_5주차_9회차_미니퀴즈.docx"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("미니퀴즈 (주관식) — 트리 & 수식")
    set_kr(run, size=14, bold=True)

    # 부제
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"날짜: {datetime.now().strftime('%Y-%m-%d')}   이름: ____________   학번: ____________________")
    set_kr(run, size=10, italic=True)

    # 안내
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("※ 간단명료하게 쓰세요. 풀이가 필요한 경우 핵심만 적기.")
    set_kr(run, size=9)

    # 문항
    q = 1
    add_short(doc, q, "트리(Tree)의 정의를 한 줄로 쓰시오", lines=2); q+=1
    add_short(doc, q, "이진 트리(Binary Tree)의 정의를 한 줄로 쓰시오.", lines=2); q+=1
    add_short(doc, q, "이진 트리 노드가 가지는 세 가지 기본 정보를 쓰시오.", lines=2); q+=1
    add_short(doc, q, "다음 수식을 이진 트리로 나타낼 때, 모양이 달라지는 이유를 한 줄로 설명하시오. 1) 1 + 2 + 3, 2) 1 + 2 * 3.", lines=2); q+=1
    add_short(doc, q, "수식 트리에서 EVAL(node) 함수가 하는 일을 간단히 쓰시오.", lines=2); q+=1
    add_short(doc, q, "트리 순회 방식 중 Preorder, Inorder, Postorder의 공통점과 차이를 한 줄씩 요약하시오.", lines=2); q+=1

    # ★ 설문 문항 추가
    doc.add_paragraph()  # 빈 줄
    run = doc.add_paragraph().add_run("※ 수업 피드백")
    set_kr(run, size=11, bold=True)

    add_short(doc, "S1", "오늘 수업 분량은 적절했습니까? (너무 적다 / 적절하다 / 많다 중 선택하고 간단히 이유)", lines=2)
    add_short(doc, "S2", "추가적으로 원하는 설명이나 개선 사항이 있으면 간단히 적어주세요.", lines=2)

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    path = build_one_page_quiz()
    print("Saved:", path)
