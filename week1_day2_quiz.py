# pip install python-docx   # 필요 시 활성화
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

TITLE = "자료구조 2회차 미니 퀴즈"
OUTPUT_PATH = "자료구조_2회차_미니퀴즈.docx"

# ----------------------------
# 콘텐츠 정의 (학생용)
# ----------------------------
OX_QUESTIONS = [
    "들여쓰기는 스페이스 4칸을 권장한다.  [ ]",
    "switch의 case 라벨은 switch와 같은 열(column)에 맞춘다.  [ ]",
    "함수 정의의 여는 { 는 함수 헤더와 같은 줄에 둔다.  [ ]",
    "sizeof 뒤에는 공백을 넣지 않으며, 괄호 안쪽에도 공백을 두지 않는다.  [ ]",
    ". 와 -> 주변에는 공백을 두지 않는다.  [ ]",
    "포인터 선언에서 * 는 변수 이름 쪽에 붙인다(예: char *p;).  [ ]",
    "if/else가 모두 단문이면 중괄호를 생략할 수 있다.  [ ]",
    "if/else 중 한쪽만 복합문이면 양쪽 모두 중괄호를 쓴다.  [ ]",
    "_Bool은 0/1만 저장하며, true/false를 쓰려면 <stdbool.h>가 필요하다.  [ ]",
    "전역 변수는 편의를 위해 자주 사용하는 것이 좋다.  [ ]",
]

MC_QUESTIONS = [
    # (문항, 보기 리스트)
    ("올바른 포인터 선언(커널 스타일)은?",
     ["A) char* buf;", "B) char *buf;", "C) char * buf ;", "D) char  *buf"]),
    ("배열 리스트 삭제 시, 오른쪽 원소를 한 칸씩 앞으로 옮기는 루프는? (크기: size)",
     ["A) for (i = k; i <= size; ++i) A[i] = A[i + 1];",
      "B) for (i = k; i <  size; ++i) A[i] = A[i + 1];",
      "C) for (i = k; i <  size - 1; ++i) A[i] = A[i];"]),
    ("삼항 연산자 공백이 올바른 것은?",
     ["A) x=(a>b)?a:b;", "B) x = (a > b) ? a : b;", "C) x=(a > b)?a : b ;"]),
    ("에러 경로에서 goto 라벨 이름으로 가장 좋은 것은?",
     ["A) error1: / error2:", "B) out_free_buf: / err_free_bar: / err_free_foo:", "C) end:", "D) fail:"]),
]

CODE_Q1 = (
    "int x = 0;\n"
    "void f(void)\n"
    "{\n"
    "        int a = x;          /* 전역 x = 0 */\n"
    "        int x = 1;          /* 함수 지역 x */\n"
    "        for (int x = 2; x < 4; ++x) {\n"
    "                int b = x;  /* 블록 x */\n"
    "        }\n"
    "        a = x;              /* 여기서 a = ? */\n"
    "}\n"
)

CODE_Q2 = (
    "#include <stdio.h>\n\n"
    "void set3(int *p)\n"
    "{\n"
    "        p[0] = 1; p[1] = 2; p[2] = 3;\n"
    "}\n\n"
    "int main(void)\n"
    "{\n"
    "        int a[3] = {0};\n"
    "        set3(a);\n"
    "        for (int i = 0; i < 3; ++i)\n"
    "                printf(\"%d \", a[i]);\n"
    "        printf(\"\\n\");\n"
    "        return 0;\n"
    "}\n"
)

CODE_WRITE_TASKS = [
    ("arraylist_initialize() 코드 적어보기",
     "void arraylist_initialize(struct ArrayList *A, int cap)\n"
     "{\n"
     "        /* TODO: initialize A->data(cap+1), A->size=0, A->capacity=cap */\n"
     "}\n"),

    ("arraylist_insert() 코드 적어보기",
     "void arraylist_insert(struct ArrayList *A, int k, int val)\n"
     "{\n"
     "        /* TODO: shift [k..size] to right, put val at k, size++ */\n"
     "}\n"),

    ("arraylist_delete() 코드 적어보기",
     "void arraylist_delete(struct ArrayList *A, int k)\n"
     "{\n"
     "        /* TODO: shift [k+1..size] to left, size-- */\n"
     "}\n"),

    ("arraylist_retrieve() 코드 적어보기",
     "int arraylist_retrieve(struct ArrayList *A, int k)\n"
     "{\n"
     "        /* TODO: return A->data[k] */\n"
     "        return 0;\n"
     "}\n"),

    ("arraylist_insert_var() 코드 적어보기",
     "void arraylist_insert_var(struct ArrayList *A, int k, int val)\n"
     "{\n"
     "        /* TODO: if full then resize, then call insert */\n"
     "}\n"),

    ("termlist_initialize() 코드 적어보기",
     "void termlist_initialize(struct TermList *A, int cap)\n"
     "{\n"
     "        /* TODO: alloc coeff/exp (cap+1), size=0, capacity=cap */\n"
     "}\n"),

    ("termlist_insert() 코드 적어보기",
     "void termlist_insert(struct TermList *A, int k, int new_coeff, int new_exp)\n"
     "{\n"
     "        /* TODO: shift, write coeff/exp, size++ */\n"
     "}\n"),

    ("termlist_add() 코드 적어보기",
     "struct TermList *termlist_add(struct TermList *X, struct TermList *Y)\n"
     "{\n"
     "        /* TODO: merge by exponent desc; sum coeffs when exponents equal */\n"
     "        return NULL;\n"
     "}\n"),
]

# ----------------------------
# 정답/해설 (교사용)
# ----------------------------
ANSWERS_OX = ["X", "O", "X", "O", "O", "O", "O", "O", "O", "X"]
ANSWERS_MC = [ "B", "B", "B", "B" ]
ANSWERS_CODE = {
    "스코프 문제": "1",
    "배열/포인터 출력": "1 2 3 (그리고 줄바꿈)"
}

# ----------------------------
# 서식 유틸
# ----------------------------
def set_korean_base_font(doc, font_name="Malgun Gothic", size_pt=11):
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(size_pt)
    # 한글 글꼴 지정
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_numbered_list(doc, items, start_index=1):
    for i, item in enumerate(items, start=start_index):
        doc.add_paragraph(f"{i}) {item}", style="List Number")

def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def add_code_block(doc, code_text, font_name="Consolas", font_size=10):
    """
    간단한 코드 블록(모노스페이스) 추가.
    """
    for line in code_text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        # 한글 코드폰트 동시 지정 (필요 시)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def set_kr(run, size=11, bold=False, italic=False, font_name="Malgun Gothic"):
    try:
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

# ----------------------------
# 본문 작성
# ----------------------------
def build_quiz_docx(output_path=OUTPUT_PATH):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("미니퀴즈 (주관식) — 배열 리스트 & 다항식")
    set_kr(run, size=14, bold=True)

    # 부제
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"날짜: {datetime.now().strftime('%Y-%m-%d')}   이름: ____________   학번: ____________________")
    set_kr(run, size=10, italic=True)

    # 안내
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("※ 간단명료하게 쓰세요. 풀이가 필요한 경우 핵심만 적기.")
    set_kr(run, size=9)

    for i, (name, code_skeleton) in enumerate(CODE_WRITE_TASKS, start=1):
        add_paragraph(doc, f"{i}) {name}", bold=True)
        add_code_block(doc, code_skeleton)
        # 답안 공간 (빈 줄 몇 개)
        for _ in range(3):
            doc.add_paragraph()

    add_numbered_list(doc, OX_QUESTIONS, start_index=1)

    # B. 객관식
    qnum = 11
    for stem, choices in MC_QUESTIONS:
        add_paragraph(doc, f"{qnum}) {stem}")
        add_bullet_list(doc, choices)
        qnum += 1

    # C. 코드 출력

    add_paragraph(doc, "15) 다음에서 함수 끝나기 직전 a의 값은?")
    add_code_block(doc, CODE_Q1)

    add_paragraph(doc, "16) 다음 프로그램의 출력 결과(한 줄)는?")
    add_code_block(doc, CODE_Q2)

    # 저장
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    build_quiz_docx()
