# -*- coding: utf-8 -*-
# 프로그래밍 과제 DOCX 생성 스크립트 + 수업 분량 설문
# 필요: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def set_kr(run, size=11, bold=False, italic=False, font_name="Malgun Gothic"):
    try:
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_line(doc, width_chars=60):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("─" * width_chars)
    set_kr(run, size=9)

def add_short(doc, num, stem, lines=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"{num}. {stem}")
    set_kr(run, size=11)
    for _ in range(lines):
        add_line(doc)

def build_one_page_quiz(output_path="doc/자료구조_과제1_리스트.docx"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("과제 — 배열리스트 & 링크드 리스트 & 다항식")
    set_kr(run, size=14, bold=True)

    # 부제
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"날짜: {datetime.now().strftime('%Y-%m-%d')}   이름: ____________   학번: ____________________")
    set_kr(run, size=10, italic=True)

    # 안내
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("※ 간단명료하게 쓰세요. 풀이가 필요한 경우 핵심만 적기.")
    set_kr(run, size=9)

    # 문항
    q = 1
    add_short(doc, q, "리스트의 기본 연산 세 가지(INSERT/DELETE/RETRIEVE)를 쓰고 각 의미를 한 줄씩 요약하시오.", lines=0); q+=1
    add_short(doc, q, "고정크기 배열 리스트 INSERT(A, k, val)의 ‘조건’과 핵심 동작을 간단히 쓰시오.", lines=0); q+=1
    add_short(doc, q, "고정크기 배열 리스트 INSERT(A, k, val)의 코드를 작성하시오.", lines=0); q+=1
    add_short(doc, q, "단방향 링크드 리스트 INSERT(A, k, val)의 ‘조건’과 핵심 동작을 간단히 쓰시오.", lines=0); q+=1
    add_short(doc, q, "단방향 링크드 리스트 INSERT(A, k, val)의 코드를 작성하시오.", lines=0); q+=1
    add_short(doc, q, "고정크기 배열 리스트에서 INSERT_AFTER(p, x)을 구현하는 방법을 쓰시오.", lines=0); q+=1
    add_short(doc, q, "고정크기 배열 리스트에서 INSERT_AFTER(p, x) 코드를 작성하시오.", lines=0); q+=1
    add_short(doc, q, "다항식 ADD의 핵심 아이디어를 쓰시오.", lines=0); q+=1
    add_short(doc, q, "다항식 ADD(X, Y)를 고정크기 배열 리스트로 구현하시오.", lines=0); q+=1
    add_short(doc, q, "다항식 ADD(X, Y)를 단방향 링크드 리스트로 구현하시오.", lines=0); q+=1
    add_short(doc, q, "양방향 링크드 리스트에서 DLIST_INSERT_FIRST(l, node)를 설명하시오.", lines=0); q+=1
    add_short(doc, q, "양방향 링크드 리스트에서 DLIST_INSERT_FIRST(l, node)의 코드를 작성하시오.", lines=0); q+=1
    add_short(doc, q, "양방향 링크드 리스트에서 DLIST_INSERT_LAST(l, node)를 설명하시오.", lines=0); q+=1
    add_short(doc, q, "양방향 링크드 리스트에서 DLIST_INSERT_LAST(l, node)의 코드를 작성하시오.", lines=0); q+=1
    add_short(doc, q, "양방향 링크드 리스트에서 DLIST_INSERT_AFTER(l, prev, node)를 설명하시오.", lines=0); q+=1
    add_short(doc, q, "양방향 링크드 리스트에서 DLIST_INSERT_AFTER(l, prev, node)의 코드를 작성하시오.", lines=0); q+=1

    # ★ 설문 문항 추가
    doc.add_paragraph()  # 빈 줄
    run = doc.add_paragraph().add_run("제출 기한: 2025-09-26 수업시간")
    set_kr(run, size=11, bold=True)

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    path = build_one_page_quiz()
    print("Saved:", path)
